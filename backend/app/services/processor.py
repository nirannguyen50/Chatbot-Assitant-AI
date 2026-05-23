"""
Document processor: parse PDF/Word/Excel/txt thành text, rồi chunk thành các đoạn nhỏ.
Hỗ trợ: .pdf, .docx, .xlsx, .xls, .csv, .txt
"""
from pathlib import Path
from typing import List

import pdfplumber
from docx import Document as DocxDocument


def extract_text(file_path: str, mime_type: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if mime_type == "application/pdf" or suffix == ".pdf":
        return _extract_pdf(file_path)

    if suffix == ".docx" or "wordprocessingml" in mime_type:
        return _extract_docx(file_path)

    if suffix in (".doc",) or mime_type == "application/msword":
        raise ValueError("File .doc cũ chưa được hỗ trợ. Vui lòng chuyển sang .docx")

    if suffix in (".xlsx", ".xls") or "spreadsheetml" in mime_type or "excel" in mime_type.lower():
        return _extract_excel(file_path)

    if suffix == ".csv" or mime_type == "text/csv":
        return _extract_csv(file_path)

    if suffix == ".txt" or mime_type.startswith("text/"):
        return path.read_text(encoding="utf-8", errors="replace")

    raise ValueError(f"Định dạng không được hỗ trợ: {suffix}. Hỗ trợ: PDF, DOCX, XLSX, XLS, CSV, TXT")


def _extract_excel(file_path: str) -> str:
    """Đọc tất cả các sheet trong file Excel, chuyển thành text dạng bảng."""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, data_only=True)
    sections = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Bỏ qua hàng toàn None
            cells = [str(c).strip() if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sections.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
    return "\n\n".join(sections)


def _extract_csv(file_path: str) -> str:
    """Đọc CSV, mỗi hàng thành 1 dòng text."""
    import csv
    rows = []
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(file_path, encoding=enc, newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    if any(cell.strip() for cell in row):
                        rows.append(" | ".join(cell.strip() for cell in row))
            break
        except UnicodeDecodeError:
            continue
    return "\n".join(rows)


def _extract_pdf(file_path: str) -> str:
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    return "\n\n".join(pages)


def _extract_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Thêm text từ các bảng
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def chunk_text(text: str, chunk_size: int = 400, overlap: int = 60) -> List[str]:
    """
    Chia text thành các chunk theo số từ, có overlap để không mất context.
    chunk_size=400 từ ≈ 500-600 tokens, phù hợp với embedding model.
    """
    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0

    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])

        # Chỉ giữ chunk có đủ nội dung
        if len(chunk.strip()) > 80:
            chunks.append(chunk)

        if end == len(words):
            break
        start += chunk_size - overlap

    return chunks
